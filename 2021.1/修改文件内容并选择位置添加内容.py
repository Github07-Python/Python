#-*- encoding:utf_8 -*-
from docx import Document
from docx.shared import Inches
#from log
import logger
import os

'''
python docx 批量修改docx文件内容和选择位置添加内容的初步尝试
不喜勿喷♪(^∇^*)
'''

def folder_filelist(pathname):
    '''
    读取所有文件名，为了实现对文件夹的批处理
    '''
    filename_list = list()
    full_pathname = os.path.abspath(pathname)
    if os.path.isfile(full_pathname):
        if _is_legal_postfix(full_pathname):
            filename_list.append(full_pathname)
        else:
            raise TypeError('文件 {} 后缀名不合法！仅支持如下文件类型：{}。'.format(pathname, '、'.join(self._handle_postfix)))
    elif os.path.isdir(full_pathname):
        for relpath, _, files in os.walk(full_pathname):
            for name in files:
                filename = os.path.join(full_pathname, relpath, name)
                if _is_legal_postfix(filename):
                    filename_list.append(os.path.join(filename))
        return filename_list
    else:
        raise TypeError('F:\\a {} 不存在或不合法！'.format(pathname))

def _is_legal_postfix(filename):
    _handle_postfix = ['doc', 'docx', 'ppt', 'pptx', 'xls', 'xlsx', 'pdf']
    return filename.split('.')[-1].lower() in _handle_postfix and not os.path.basename(filename).startswith(
        '~')


def insert_dash(string, index, insert_content):
    '''
    str中插入字符
    '''
    return string[:index] + insert_content + string[index:]

def find_replace(docxname, findname, replacename):
    '''
    docx中替换内容
    '''
    document = Document(docxname)

    for paragraph in document.paragraphs:
        line = paragraph.text
        logger.info(line)
        paragraph.clear()
        if line.find(findname) != -1:
            line = insert_dash(line, line.find(findname), replacename)
            line = line.replace(findname,"")
        paragraph.add_run(line)
        logger.info(line)

    document.save("1" + docxname)

def find_add_paragraph(docxname, findname, insert_paragraph):
    '''
    docx中想要的位置，插入段落
    '''
    document = Document(docxname)

    for paragraph in document.paragraphs:
        line = paragraph.text
        logger.info(line)
        if line.find(findname) != -1:
            for content in insert_paragraph:
                paragraph.add_run('\n'+content)
                logger.info(content)

    document.save("1" + docxname)

if __name__ == '__main__':

    for filename in folder_filelist('F:\\a'):
        logger.info(filename)
    find_replace('1.docx', '666', '6666')

    find_add_paragraph('1.docx', '6', ['你好', 'spa'])

#原文链接：https://blog.csdn.net/zw17302560727/article/details/100929891